"""Extract text from DOCX (and attempt DOC via libreoffice)."""
from __future__ import annotations

import logging
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_docx(file_path: str | Path) -> str:
    """Extract text from .docx files preserving paragraph order."""
    from docx import Document
    doc = Document(str(file_path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Include tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def extract_doc(file_path: str | Path) -> str:
    """Convert .doc → .docx via LibreOffice, then extract."""
    file_path = Path(file_path)
    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", tmpdir, str(file_path)],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed for {file_path}: {result.stderr}"
            )
        converted = Path(tmpdir) / (file_path.stem + ".docx")
        if not converted.exists():
            raise RuntimeError(f"Converted file not found at {converted}")
        return extract_docx(converted)
